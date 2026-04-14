from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from src.parsers.common.models import DocumentClassification, DocumentElement, ParsedDocument, ParsedPage
from src.parsers.common.serialization import build_document_summary, extract_markdown_sections
from src.shared.io import iso_now, make_artifact_stem


class DocxParser:
    parser_name = "python-docx-parser"

    def parse(self, path: Path, classification: DocumentClassification) -> ParsedDocument:
        document = DocxDocument(path)
        elements: list[DocumentElement] = []
        markdown_parts: list[str] = []
        order = 1

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name if paragraph.style else ""
            lower_style = style_name.lower()
            if lower_style.startswith("heading"):
                heading_level = self._extract_heading_level(style_name)
                markdown_parts.append(f"{'#' * heading_level} {text}")
                element_type = "heading"
            else:
                markdown_parts.append(text)
                element_type = "text"

            elements.append(
                DocumentElement(
                    element_id=f"p1-e{order}",
                    element_type=element_type,
                    page_number=1,
                    order=order,
                    text=text,
                    metadata={"style": style_name},
                )
            )
            order += 1

        for table in document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            markdown = self._rows_to_markdown(rows)
            if not markdown:
                continue
            markdown_parts.append(markdown)
            elements.append(
                DocumentElement(
                    element_id=f"p1-e{order}",
                    element_type="table",
                    page_number=1,
                    order=order,
                    markdown=markdown,
                    metadata={
                        "row_count": len(rows),
                        "column_count": max((len(row) for row in rows), default=0),
                        "cells": rows,
                    },
                )
            )
            order += 1

        markdown = "\n\n".join(markdown_parts)
        parsed = ParsedDocument(
            document_id=make_artifact_stem(path),
            source_name=path.name,
            source_path=path.as_posix(),
            extension=path.suffix.lower(),
            status="parsed",
            parser_name=self.parser_name,
            classification=classification,
            metadata={"paragraph_count": len(document.paragraphs), "table_count": len(document.tables)},
            sections=extract_markdown_sections(markdown),
            pages=[
                ParsedPage(
                    page_number=1,
                    width=0.0,
                    height=0.0,
                    text_length=sum(len(part) for part in markdown_parts),
                    elements=elements,
                    metadata={"synthetic_page": True},
                )
            ],
            markdown=markdown,
            created_at=iso_now(),
        )
        parsed.summary = build_document_summary(parsed)
        return parsed

    def _extract_heading_level(self, style_name: str) -> int:
        digits = "".join(char for char in style_name if char.isdigit())
        if digits:
            return min(max(int(digits), 1), 6)
        return 1

    def _rows_to_markdown(self, rows: list[list[str]]) -> str:
        if not rows:
            return ""
        width = max((len(row) for row in rows), default=0)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        total_rows = len(normalized)
        header = normalized[0]
        separator = ["---"] * width
        body = normalized[1:] or [[""] * width]
        markdown_rows = [header, separator, *body]
        table_meta = f"[표: {total_rows}행 × {width}열]"
        return table_meta + "\n" + "\n".join("| " + " | ".join(row) + " |" for row in markdown_rows)
